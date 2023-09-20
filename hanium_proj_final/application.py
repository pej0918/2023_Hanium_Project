from flask import Flask, render_template, request, flash, redirect, url_for, session
import sys

# 점자 번역
import louis
from KorToBraille.KorToBraille import KorToBraille
from BrailleToKorean.BrailleToKor import BrailleToKor

# 문서 요약: 모델 및 Tokenizer 불러오기
import torch 
import transformers  
 
     
from transformers import PreTrainedTokenizerFast, BartForConditionalGeneration

#워드파일
# 가장 기본적인 기능(문서 열기, 저장, 글자 쓰기 등등)
import docx
from docx.oxml.ns import qn

# 문단 정렬
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 문자 스타일 변경
from docx.enum.style import WD_STYLE_TYPE

from docx2pdf import convert

# 영어 철자 교정
from autocorrect import Speller

# TextToSpeech
from gtts import gTTS
import IPython.display as ipd

# 번역
import requests
import json
import warnings

#PDF에서 글자 읽어오기 및 PDF로 저장
warnings.filterwarnings('ignore')
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
import io
import re
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

import aspose.pdf as ap


application = Flask(__name__)


application.secret_key = '1000000000'

#캐시 제어
application.config["CACHE_TYPE"] = "null"
@application.after_request
def set_response_headers(response):
    response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    return response

#Index 페이지
@application.route("/")
def index():
    return render_template("index.html")
    #return redirect(url_for("index"))
    
#Guide 페이지
@application.route("/Guide")
def guide():
    return render_template("Guide.html")

#번역 페이지
@application.route("/Translate", methods=['GET','POST'])
def translateSelect():
    file = 'static/result_pdf/result.docx'
    if os.path.isfile(file):
        os.remove(file)
        print('result 파일 삭제')
        
        
    file = 'static/upload_pdf/upload.pdf'
    if os.path.isfile(file):
        os.remove(file)
        print('upload 파일 삭제')
        
    if request.method == 'POST':
        input_language = request.form['input_check']
        output_language = request.form['output_check']
        grammar_check = request.form.get('grammar_check')

        # 세션에 데이터 저장
        session['input_language'] = input_language
        session['output_language'] = output_language
        session['grammar_check'] = grammar_check

        print("입력 언어: " + str(input_language))
        print("출력 언어: " + str(output_language))

        if grammar_check:
            print("맞춤법 검사 O")
        else:
            print("맞춤법 검사 X")

        return redirect("getInput")
    return render_template("Translate.html")


#----------------------------------------기능-------------------
#한글 맞춤법 검사
def grammar_test_ko(text):
    data=None
    response=None
    origin_text=None
    correct_text=None
    notcorrect=None
    change_list=[]
    res='잘못된 내용이 없습니다.'
    if request.method == 'POST':
        # 맞춤법 검사 요청 (requests) -> API요청
        response = requests.post('http://164.125.7.61/speller/results', data={'text1': text})

        # 응답에서 필요한 내용 추출 (html 파싱)
        data = response.text.split('data = [', 1)[-1].rsplit('];', 1)[0]
        
        try:
            data = json.loads(data)
            # 교정 내용 담기
            for i in range(len(data['errInfo'])):
                correct_text=data['errInfo'][i]['candWord'].split('|')[0]
                if correct_text == "":
                    return "grammar_error"
                notcorrect=data['errInfo'][i]['orgStr']
                print("correct: " + correct_text)
                print("not correct: " + notcorrect)

                # 기존 내용 -> 교정 내용 change
                text = text.replace(notcorrect,correct_text)
                origin_text = notcorrect+' → '+correct_text
                print("origin_text: " + origin_text)
                print("text: " + text)
                change_list.append(origin_text)

            change_list.append(text)
            return change_list
        except:
            change_list.append(res)
            change_list.append(text)
            return change_list

        
#영어 맞춤법 검사
def grammar_test_en(text):
    data=None
    response=None
    origin_text=None
    correct_text=None
    notcorrect=None
    change_list=[]
    res='잘못된 내용이 없습니다.'
    #spell=Speller('en')
    if request.method == 'POST':
        #영어 철자 교정 : beacuse i love you. -> because i love you.
        #text=str(spell(text))
        #text = translateLanguage(text, 'en', 'ko')
        # 맞춤법 검사 요청 (requests) -> API요청
        response = requests.post('http://164.125.7.61/speller/results', data={'text1': text})
        # 응답에서 필요한 내용 추출 (html 파싱)
        data = response.text.split('data = [', 1)[-1].rsplit('];', 1)[0]
        # 파이썬 딕셔너리 형식으로 변환
        try:
            data = json.loads(data)
            # 교정 내용 담기
            for i in range(len(data['errInfo'])):
                correct_text=data['errInfo'][i]['candWord'].split('|')[0]
                if correct_text == "":
                    return "grammar_error"
                notcorrect=data['errInfo'][i]['orgStr']
                # 기존 내용 -> 교정 내용 change
                text=text.replace(notcorrect,correct_text)
                text = translateLanguage(text,'ko','en')
                print(text)
                origin_text=notcorrect + '->' + correct_text
                change_list.append(origin_text)
                print(change_list)
            change_list.append(text)
            return change_list
        except:
            #text = translateLanguage(text,'ko','en')
            change_list.append(res)
            change_list.append(text)
            return change_list



# 번역
papago_id = "ucWN1NYU3VpbYcuxAKTq"
papago_pw = "shWp9ApkCt"

def translateLanguage(text, lan1, lan2):
    trans_data = None
    data = {'text': text, 'source': lan1, 'target': lan2}
    url = "https://openapi.naver.com/v1/papago/n2mt"
    global papago_id
    global papago_pw

    header = {"X-Naver-Client-Id": papago_id,
              "X-Naver-Client-Secret": papago_pw}

    response = requests.post(url, headers=header, data=data)
    rescode = response.status_code

    if(rescode == 200):
        send_data = response.json()
        trans_data = (send_data['message']['result']['translatedText'])
    else:
        print("Error Code:", rescode)
        return -10000
    return trans_data

# 한글 음성 변환
def tts_ko(text):
    tts = gTTS(text=text, lang='ko')
    audio_file_path = "static/audio/result_kor.mp3"
    tts.save(audio_file_path)
    return audio_file_path


def tts_en(text):
    tts = gTTS(text=text, lang='en')
    audio_file_path = "static/audio/result_eng.mp3"
    tts.save(audio_file_path)
    return audio_file_path


#pdf에서 글자 읽어오기
def convert_pdf_to_txt(filepath):
    rsrcmgr = PDFResourceManager()
    retstr = io.StringIO()
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, laparams=laparams)
    fp = open(filepath, 'rb')
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    password = ""
    maxpages = 0
    caching = True
    pagenos = set()
    for page in PDFPage.get_pages(fp, pagenos, maxpages=maxpages,
                                  password=password,
                                  caching=caching,
                                  check_extractable=True):
        interpreter.process_page(page)
    text = retstr.getvalue()

    fp.close()
    device.close()
    retstr.close()
    text=re.sub('\n','',text)[:-1]
    text=re.sub('Evaluation Only. Created with Aspose.PDF. Copyright 2002-2023 Aspose Pty Ltd.','',text)
    text=re.sub('  ', ' ', text)
    return text

#----------------------------------------기능 끝


# 입력 페이지
@application.route("/getInput", methods=['GET', 'POST'])
def getInput():
    
    result=None
    text_input=None
    input_language=None
    output_language=None
    grammar_check=None
    audio_file_path = None
    pdf_file_path=None
    result2=None
    change_result=[]
    total_result=None
    grammar_result=None
    grammar_before_result = None
    
    if request.method == 'POST':
        if 'text_input' in request.form:
            text_input = request.form['text_input']
            

            # 세션 데이터 가져오기
            input_language = session.get('input_language')
            output_language = session.get('output_language')
            grammar_check = session.get('grammar_check')
            result = text_input
            
            if text_input == "":
                flash("텍스트를 입력해주세요!")
                return render_template("getInput.html")
            # 입력 텍스트가 한글 또는 한글 점자이면
            if input_language == 'kor' or input_language == 'korb':
                if input_language == 'korb':
                    result = BrailleToKor().translation(result) # 한글 점자 번역
                    grammar_before_result = result
                if grammar_check == "grammar":
                    #result2=result
                    result = grammar_test_ko(str(result)) # 한글 맞춤법 검사
                    print("################################\n")
                    print("바뀐 텍스트: "+ str(result[-1]))
                    print("################################")
                    change_result=result[:-1]
                    result=str(result[-1])
                    grammar_result = result
                    
                    print(grammar_result)
                    #change_result= grammar_test_ko(result2)
                    change_result="<br> ▪ ".join(change_result)
                    print(change_result)
                    pattern = re.compile('[^a-zA-Z가-힣\s]')
                    cleaned_text = re.sub(pattern, '', result)
                    if cleaned_text == "":
                        flash("텍스트를 확인해주세요!")
                        return render_template("getInput.html")
                    print('맞춤법 비교 결과 GET')
                # 출력이 한글 점자이면
                if output_language == 'korb':
                    result = KorToBraille().korTranslate(result)
                # 출력이 한글 문자이면
                elif output_language == 'kor':
                    print("음성 링크 반환")
                    audio_file_path = tts_ko(result) # 음성 변환
                # 출력이 영어 문자 or 영어 점자이면
                else:
                    result = translateLanguage(result, 'ko', 'en')
                    if result == -10000:
                        flash("텍스트를 다시 확인해주세요!")
                        return render_template("getInput.html")
                    if output_language == 'eng':
                        audio_file_path = tts_en(result) # 음성 변환
                    elif output_language == 'engb':
                        result = louis.translateString(["braille-patterns.cti", "en-us-g2.ctb"], result)
            # 입력 텍스트가 영어 문자 또는 영어 점자이면
            else:
                if input_language == 'engb':
                    result =  louis.backTranslateString(["braille-patterns.cti", "en-us-g2.ctb"], result) # 영어 점자 번역
                    grammar_before_result = result
                    print("문법교정 전"+grammar_before_result)
                if grammar_check:
                    #result2=result
                    result = grammar_test_en(result)# 영어 맞춤법 검사
                    print(result)
                    change_result=result[:-1]
                    result=str(result[-1])
                    grammar_result = result
                    #change_result=grammar_test_en(result2)
                    change_result="<br> ▪ ".join(change_result)
                    print(change_result)
                    pattern = re.compile('[^a-zA-Z가-힣\s]')
                    cleaned_text = re.sub(pattern, '', result)
                    
                    print("cleaned_text: " + cleaned_text)
                    if cleaned_text == "":
                        flash("텍스트를 확인해주세요!")
                        return render_template("getInput.html")
                    print('맞춤법 비교 결과 GET')

                
                # 출력이 영어 문자이면
                if output_language == 'eng':
                    audio_file_path = tts_en(result) # 음성 변환
                # 출력이 영어 점자이면
                elif output_language == 'engb':
                    result = louis.translateString(["braille-patterns.cti", "en-us-g2.ctb"], result)
 
                # 출력이 한글 문자 or 한글 점자이면
                else:    
                    result = translateLanguage(result, 'ko', 'en')
                    if result == -10000:
                        flash("텍스트를 다시 확인해주세요!")
                        return render_template("getInput.html")

                    result = translateLanguage(result, 'en', 'ko') # 한글로 번역
                    if result == -10000:
                        flash("텍스트를 다시 확인해주세요!")
                        return render_template("getInput.html")
                    if output_language == 'kor':
                        audio_file_path = tts_ko(result) # 음성 변환
                    elif output_language == 'korb':
                        result = KorToBraille().korTranslate(result)
        
        print("grammar_before_result:" + grammar_before_result)
        
        return redirect(url_for("result", result=result, total_result=result, grammar_result=grammar_result, grammar_before_result=grammar_before_result, compare_result=change_result, pdf_file_path=pdf_file_path, audio_file_path=audio_file_path, input_language=input_language, output_language=output_language, text_input=text_input))
    return render_template("getInput.html")



#UPLOAD_FOLDER = os.getcwd() +'/upload' # 절대 파일 경로
UPLOAD_FOLDER = 'static/upload_pdf'
#DOWN_FOLDER=os.getcwd()+'/result'
DOWN_FOLDER = 'static/result_pdf'
ALLOWED_EXTENSIONS = set(['pdf'])

UPLOAD_FILE = None
UPLOAD_FILE_TRANS=None



# 결과 페이지
@application.route("/result", methods=['GET', 'POST'])
def result():
    input_language = request.args.get('input_language')
    output_language = request.args.get('output_language')
    result_value = request.args.get('result')
    text_input = request.args.get('text_input')
    audio_file_path = request.args.get('audio_file_path')
    filepathtosave = request.args.get('pdf_file_path')
    total_result=request.args.get('total_result')
    change_result=request.args.get('compare_result')
    grammar_result=request.args.get('grammar_result')
    grammar_before_result = request.args.get('grammar_before_result')
    
    if change_result != None:
        change_result = "▪ " + change_result
    
    print("result 페이지 실행")


    '''
    filename=re.sub(r"[^\uAC00-\uD7A30-9a-zA-Z\s]", "d",result_value[:10])
    filename = filename+'.pdf'
    filepathtosave = os.path.join(DOWN_FOLDER,filename)
    filepathtosave=filepathtosave.replace('.pdf','.docx')
    '''
    filepathtosave = 'static/result_pdf/result.docx'

    #pdf=canvas.Canvas(filepathtosave)
    #pdf.drawString(50, 800, result_value)
    #print(result_value)
    #pdf.save()
    #outputFile=filepathtosave

    doc = docx.Document()
    para = doc.add_paragraph()
    run = para.add_run(result_value)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    doc.save(filepathtosave)

    print(result_value)
    #doc.add_paragraph(result_value)

    #convert(filepathtosave, outputFile)
    print('업로드 완료')
    print(filepathtosave)
    return render_template("result_forcheck.html", result=result_value, total_result=result, grammar_before_result=grammar_before_result, grammar_result=grammar_result, compare_result=change_result, pdf_file_path=filepathtosave, audio_file_path=audio_file_path, input_language=input_language, output_language=output_language, text_input=text_input)




'''

@application.route("/compare", methods=['GET', 'POST'])
def compare():
    print("비교 시작")
    compare_result=None
    result = None
    
    #if request.method == 'POST':
    text_input = request.form.get('text_input')
    text_input = str(text_input)
    language_input = request.form.get('input_language')
    language_input = str(language_input)
        
    #입력 언어가 한국이면 
    if language_input=='kor' or language_input=='korb':
        print("한글")
        result = str(grammar_test_ko(text_input))
        print(result)
        compare_result = result[:-1]
        print(compare_result)
        result = result[-1]
    else: #영어이면
        print("영어")
        result = str(grammar_test_en(text_input))
        compare_result = result[:-1]
        result = result[-1]
    
    return render_template("compare.html", text_input=text_input, result=result,compare_result=compare_result)


'''                    



@application.route("/Summary_choose2", methods=['GET','POST'])
def summary_main2():
    return render_template("Summary_choose2.html")



#PDF 업로드 및 다운로드
import os
from flask import Blueprint, send_file, request, redirect, url_for, render_template

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1] in ALLOWED_EXTENSIONS


#파일 업로드
@application.route('/fileupload_summary_en', methods=['GET', 'POST'])
def upload_file_eng():
    file = 'static/upload_pdf/upload.pdf'
    if os.path.isfile(file):
        os.remove(file)
        print('파일 삭제')
    
    pdf_result=None
    input_ids=None
    result=None
    summary_text_ids=None
    pdf=None
    filepathtosave=None


    if request.method == 'POST':
        print(request.files)
        file = request.files['savefile']
        if file and allowed_file(file.filename):
            '''
            filename = file.filename
            print(filename)
            filepathtosave = os.path.join(UPLOAD_FOLDER, filename)
            
            '''
            filepathtosave = 'static/upload_pdf/upload.pdf'
            file.save(filepathtosave)
            
            print(filepathtosave)
            UPLOAD_FILE = filepathtosave
            pdf_result=convert_pdf_to_txt(filepathtosave)
            
            print("PDF 내용 출력:")
            print(pdf_result)
    
            print('업로드 완료')
            return render_template('Summary_en.html',text_result=pdf_result)
    #return redirect(url_for("result", result=result, total_result=result,compare_result=change_list, pdf_file_path=filepathtosave, audio_file_path=audio_file_path, input_language=input_language, output_language=output_language, text_input=text_input))
    return render_template('Summary_en.html',text_result=pdf_result)



#파일 업로드
@application.route('/fileupload_summary', methods=['GET', 'POST'])
def upload_file_kor():
    pdf_result=None
    input_ids=None
    result=None
    summary_text_ids=None
    pdf=None
    filepathtosave=None


    if request.method == 'POST':
        print(request.files)
        file = request.files['savefile']
        if file and allowed_file(file.filename):
            filename = file.filename
            print(filename)
            filepathtosave = os.path.join(UPLOAD_FOLDER, filename)
            file.save(filepathtosave)
            print(filepathtosave)
            UPLOAD_FILE = filepathtosave
            pdf_result=convert_pdf_to_txt(filepathtosave)
            print(pdf_result)
    
            print('업로드 완료')
            return render_template('Summary.html',text_result=pdf_result)
    #return redirect(url_for("result", result=result, total_result=result,compare_result=change_list, pdf_file_path=filepathtosave, audio_file_path=audio_file_path, input_language=input_language, output_language=output_language, text_input=text_input))
    return render_template('Summary.html',text_result=pdf_result)



#문서요약
@application.route("/Summary", methods=['GET','POST'])
def Summary():
    result=None
    text=None
    pdf_file_path=None
    summary_text_ids=None
    text_length=None
    input_ids=None
    audio_file_path=None

    #뉴스 기사로 사전 학습된 모델 이용
    tokenizer=PreTrainedTokenizerFast.from_pretrained("ainize/kobart-news")
    model=BartForConditionalGeneration.from_pretrained("ainize/kobart-news")
    
    if request.method == 'POST':
        text=request.form["text_input"]
        input_ids=tokenizer.encode(text,return_tensors='pt')
        summary_text_ids = model.generate( input_ids=input_ids, bos_token_id=model.config.bos_token_id,
                                  eos_token_id=model.config.eos_token_id,
                                   length_penalty=1.0
                                   ,max_length=128 #요약문 최대 길이 설정
                                   ,min_length=32 #요약문의 최소길이 설정
                                   ,num_beams=15 #문장 생성시 다음 단어를 탐색하는 영역의 개수
                                   )
        result=tokenizer.decode(summary_text_ids[0],skip_special_tokens=True)[1:]       
        audio_file_path = tts_ko(result)
        return redirect(url_for("summary_result", result=result,  audio_file_path=audio_file_path, pdf_file_path=pdf_file_path,  text_input=text))
    return render_template("Summary.html")


# 요약 결과 페이지
@application.route("/summary_result", methods=['GET', 'POST'])
def summary_result():
    result_value = request.args.get('result')
    text_input = request.args.get('text_input')
    result_value = request.args.get('result')
    filepathtosave = request.args.get('pdf_file_path')
    audio_file_path=request.args.get('audio_file_path')
    
    print(filepathtosave)
    print("result 페이지 실행")


    if filepathtosave == None:
        filename=re.sub(r"[^\uAC00-\uD7A30-9a-zA-Z\s]", "d",result_value[:10])
        filename = filename+'.pdf'
        filepathtosave = os.path.join(DOWN_FOLDER,filename)
        #pdf=canvas.Canvas(filepathtosave)
        #pdf.drawString(50, 800, result_value)
        #print(result_value)
        #pdf.save()
        #outputFile=filepathtosave
        
        filepathtosave=filepathtosave.replace('.pdf','.docx')
        #doc = Document()
        print(result_value)
        #doc.add_paragraph(result_value)

        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run(result_value)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        doc.save(filepathtosave)
        #convert(filepathtosave, outputFile)
        print('업로드 완료')
        return render_template("Summary_result.html", result=result_value, audio_file_path=audio_file_path, pdf_file_path=filepathtosave, text_input=text_input)
    else: 
        #filepathtosave=os.path.join(UPLOAD_FOLDER,'result.pdf')
        filepathtosave=filepathtosave.replace('.pdf','.docx')
        print(filepathtosave)
        #pdf=canvas.Canvas(filepathtosave)
        #pdf.drawString(50, 800, result_value)
        #print(result_value)
        #pdf.save()
        print(result_value)
        doc = docx.Document()
        para = doc.add_paragraph()
        run = para.add_run(result_value)
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        doc.save(filepathtosave)
        #filepathtosave=filepathtosave+"_.docx"
        print('업로드 완료')
        return render_template("Summary_result.html", result=result_value,audio_file_path=audio_file_path, pdf_file_path=filepathtosave, text_input=text_input)
    

@application.route("/Summary_en", methods=['GET','POST'])
def summary_en():
    result=None
    text=None
    pdf_file_path=None
    summary_text_ids=None
    text_length=None
    input_ids=None
    audio_file_path=None
    
    #뉴스 기사로 사전 학습된 모델 이용
    tokenizer=PreTrainedTokenizerFast.from_pretrained("ainize/kobart-news")
    model=BartForConditionalGeneration.from_pretrained("ainize/kobart-news")
    
    if request.method == 'POST':
        text=request.form["text_input"]
        text=translateLanguage(text,'en','ko')
        input_ids=tokenizer.encode(text,return_tensors='pt')
        summary_text_ids = model.generate( input_ids=input_ids, bos_token_id=model.config.bos_token_id,
                                  eos_token_id=model.config.eos_token_id,
                                   length_penalty=1.0
                                   ,max_length=128 #요약문 최대 길이 설정
                                   ,min_length=32 #요약문의 최소길이 설정
                                   ,num_beams=15 #문장 생성시 다음 단어를 탐색하는 영역의 개수
                                   )
        result=tokenizer.decode(summary_text_ids[0],skip_special_tokens=True)[1:]            
        result=translateLanguage(result,'ko','en')
        audio_file_path=tts_en(result)
        return redirect(url_for("summary_result", result=result, audio_file_path=audio_file_path, pdf_file_path=pdf_file_path,  text_input=text))
    return render_template("Summary_en.html")
   


@application.route('/fileupload_translate', methods=['GET','POST'])
def upload_file_translate():
    global UPLOAD_FILE_TRANS
    result=None
    pdf=None
    audio_file_path=None
    input_language=None
    output_language=None
    text_input=None
    document=None
    text_fragment=None
    result2=None
    change_list=[]
    
    if request.method == 'POST':
        global UPLOAD_FILE_TRANS
        print(request.files)
        file = request.files['savefile']
        if file and allowed_file(file.filename):
            filename = file.filename
            filepathtosave = os.path.join(UPLOAD_FOLDER, filename)
            file.save(filepathtosave)
            print(filepathtosave)
            UPLOAD_FILE_TRANS = filepathtosave
            result=convert_pdf_to_txt(filepathtosave)
            

            # 세션 데이터 가져오기
            input_language = session.get('input_language')
            output_language = session.get('output_language')
            grammar_check = session.get('grammar_check')
            
            # 입력 텍스트가 한글 또는 한글 점자이면
            '''if input_language == 'kor' or input_language == 'korb':
                if input_language == 'korb' :
                    result = BrailleToKor().translation(result) # 한글 점자 번역
                if grammar_check:
                    result2=result
                    result = grammar_test_ko(result) # 한글 맞춤법 검사
                    change_list=grammar_test_ko(result2)[1:]
                # 출력이 한글 점자이면
                if output_language == 'korb':
                    result = KorToBraille().korTranslate(result)
                # 출력이 한글 문자이면
                elif output_language == 'kor':
                    print("음성 링크 반환")
                    audio_file_path = tts_ko(result) # 음성 변환
                # 출력이 영어 문자 or 영어 점자이면
                else:
                    result = translateLanguage(result, 'ko', 'en')
                    if output_language == 'eng':
                        audio_file_path = tts_en(result) # 음성 변환
                    elif output_language == 'engb':
                        result = louis.translateString(["braille-patterns.cti", "en-us-g2.ctb"], result)
            # 입력 텍스트가 영어 문자 또는 영어 점자이면
            else:
                if input_language == 'engb':
                    result =  louis.backTranslateString(["braille-patterns.cti", "en-us-g2.ctb"], result) # 영어 점자 번역
                if grammar_check:
                    result2=result
                    result = grammar_test_en(result)[0] # 영어 맞춤법 검사
                    change_list=grammar_test_en(result2)[1:]
                # 출력이 영어 문자이면
                if output_language == 'eng':
                    audio_file_path = tts_en(result) # 음성 변환
                # 출력이 영어 점자이면
                elif output_language == 'engb':
                    result = louis.translateString(["braille-patterns.cti", "en-us-g2.ctb"], result)
                # 출력이 영어 문자 or 영어 점자이면
                else:
                    result = translateLanguage(result, 'ko', 'en')
                    if output_language == 'en':
                        audio_file_path = tts_en(result) # 음성 변환
                    elif output_language == 'enb':
                        result = louis.translateString(["braille-patterns.cti", "en-us-g2.ctb"], result)
                    else:
                        result = translateLanguage(result, 'en', 'ko') # 한글로 번역
                        if output_language == 'kor':
                            audio_file_path = tts_ko(result) # 음성 변환
                        elif output_language == 'korb':
                            result = KorToBraille().korTranslate(result)'''
                            
            #pdfmetrics.registerFont(TTFont("맑은고딕", url_for('static', filename='malgun.ttf')))
            #pdf = canvas.Canvas(filepathtosave)
            #pdf.setFont("맑은고딕", 15)
            #pdf.drawString(50, 800, result)
            #pdf.save()
            
            filepathtosave=filepathtosave.replace('.pdf','.docx')
            doc = docx.Document()
            doc.add_paragraph(result)
            doc.save(filepathtosave)
            #convert(filepathtosave, outputFile)       
            print('업로드 완료')
            
        #return redirect(url_for("result", result=result, total_result=result,compare_result=change_list,pdf_file_path=filepathtosave, audio_file_path=audio_file_path, input_language=input_language, output_language=output_language, text_input=text_input))
        return render_template('getInput.html',text_input=result)
    #return redirect(url_for("result", result=result, total_result=result,compare_result=change_list, pdf_file_path=filepathtosave, audio_file_path=audio_file_path, input_language=input_language, output_language=output_language, text_input=text_input))
    return render_template('getInput.html',text_input=result)


if __name__ == "__main__":
    application.run(host='0.0.0.0', debug=True)
    # application.run(host='localhost', port=5000, debug=True)


